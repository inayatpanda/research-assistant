"""Phase 4.6 — Extractor tests (manuscript + uploaded file)."""
from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from research_api.db.models import (
    Article,
    ManuscriptSection,
    Project,
    new_id,
)
from research_api.services.peer_review import (
    PeerReviewExtractError,
    extract_manuscript_for_peer_review,
    extract_uploaded_document,
)


@pytest.mark.asyncio
async def test_manuscript_extractor_concatenates_sections_in_order(
    session,
) -> None:
    user = "u-mx"
    project = Project(
        id=new_id(), user_id=user, title="My Trial", study_type="RCT"
    )
    session.add(project)
    await session.flush()

    # Intentionally insert out-of-order to confirm SECTION_ORDER is honoured.
    section_html = {
        "Discussion": "<p>Discussion prose.</p>",
        "Introduction": "<p>Background <strong>matters</strong>.</p>",
        "Results": "<p>p&lt;0.001 was observed.</p>",
        "Methodology": "<p>We randomised.</p>",
        "Abstract": "<p>One-sentence abstract.</p>",
    }
    for name, html in section_html.items():
        session.add(
            ManuscriptSection(
                id=new_id(),
                user_id=user,
                project_id=project.id,
                section_name=name,
                content=html,
                word_count=len(html.split()),
            )
        )
    # Two articles → n_references == 2.
    for i in range(2):
        session.add(
            Article(
                id=new_id(),
                user_id=user,
                project_id=project.id,
                title=f"Ref {i}",
                authors=["Author A"],
            )
        )
    await session.commit()

    ex = await extract_manuscript_for_peer_review(
        project_id=project.id, user_id=user, session=session
    )

    assert ex.title == "My Trial"
    assert ex.study_type == "RCT"
    # Section headers ordered.
    abstract_idx = ex.text.index("## ABSTRACT")
    intro_idx = ex.text.index("## INTRODUCTION")
    methods_idx = ex.text.index("## METHODOLOGY")
    results_idx = ex.text.index("## RESULTS")
    discussion_idx = ex.text.index("## DISCUSSION")
    conclusion_idx = ex.text.index("## CONCLUSION")
    assert abstract_idx < intro_idx < methods_idx < results_idx < discussion_idx < conclusion_idx
    # Body text from each section is present (HTML stripped).
    assert "One-sentence abstract." in ex.text
    assert "Background matters." in ex.text
    assert "We randomised." in ex.text
    assert "p<0.001 was observed." in ex.text
    # Empty section placeholder for Conclusion.
    assert "(this section is empty)" in ex.text

    # Metadata counts.
    assert ex.metadata["n_references"] == 2
    assert ex.metadata["n_figures"] == 0
    assert ex.metadata["n_authors"] == 0


def _make_minimal_pdf_bytes(text: str = "Hello peer reviewer") -> bytes:
    """Build the smallest PDF the test environment will accept.

    We use pypdf to assemble a single blank page and lean on
    ``extract_text`` returning the empty string — the file extractor only
    needs to *not crash* on a valid PDF, and the failure-path tests
    exercise the non-PDF case.
    """
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx_bytes(text: str = "Hello peer reviewer.") -> bytes:
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.add_paragraph("A second paragraph with content.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_file_extractor_handles_pdf_and_docx() -> None:
    pdf_bytes = _make_minimal_pdf_bytes()
    pdf_ex = extract_uploaded_document(
        data=pdf_bytes, original_filename="paper.pdf"
    )
    assert pdf_ex.source_type == "uploaded_pdf"
    assert pdf_ex.mime == "application/pdf"

    docx_bytes = _make_docx_bytes()
    docx_ex = extract_uploaded_document(
        data=docx_bytes, original_filename="paper.docx"
    )
    assert docx_ex.source_type == "uploaded_docx"
    assert "wordprocessingml" in docx_ex.mime
    assert "Hello peer reviewer." in docx_ex.text
    assert "A second paragraph" in docx_ex.text


def test_file_extractor_rejects_non_document_mime() -> None:
    # Plain text byte sequence — not a PDF, not a DOCX zip.
    plain = b"Just some random text, definitely not a paper."
    with pytest.raises(PeerReviewExtractError) as exc_info:
        extract_uploaded_document(data=plain, original_filename="random.txt")
    assert exc_info.value.status_code == 415

    with pytest.raises(PeerReviewExtractError) as empty_info:
        extract_uploaded_document(data=b"", original_filename="empty")
    assert empty_info.value.status_code == 400
