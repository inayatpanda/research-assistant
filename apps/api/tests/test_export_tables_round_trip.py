"""Phase 8.7 — User-authored TipTap tables survive DOCX + PDF export."""
from __future__ import annotations

import io
from dataclasses import dataclass

import pytest
from docx import Document
from pypdf import PdfReader

from research_api.services.export.docx_export import render_docx
from research_api.services.export.pdf_export import render_pdf


@dataclass
class _Project:
    title: str = "Test"
    study_type: str = "Outcome Study"
    citation_style: str = "vancouver"


@dataclass
class _Section:
    section_name: str
    content: str


TABLE_HTML = """
<table>
  <tbody>
    <tr><th>A</th><th>B</th></tr>
    <tr><td>1</td><td>2</td></tr>
  </tbody>
</table>
"""


def test_docx_round_trip_preserves_table_cells() -> None:
    sections = [_Section("Results", TABLE_HTML)]
    data = render_docx(project=_Project(), sections=sections, bibliography=[])
    doc = Document(io.BytesIO(data))
    assert len(doc.tables) >= 1
    cells = [c.text.strip() for c in doc.tables[0].rows[0].cells]
    assert cells == ["A", "B"]
    cells2 = [c.text.strip() for c in doc.tables[0].rows[1].cells]
    assert cells2 == ["1", "2"]


def test_pdf_round_trip_preserves_table_cells() -> None:
    sections = [_Section("Results", TABLE_HTML)]
    data = render_pdf(project=_Project(), sections=sections, bibliography=[])
    reader = PdfReader(io.BytesIO(data))
    full_text = "\n".join(p.extract_text() or "" for p in reader.pages)
    for token in ("A", "B", "1", "2"):
        assert token in full_text, f"missing {token!r} in PDF text: {full_text[:200]!r}"
