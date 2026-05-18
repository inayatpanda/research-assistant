"""DOCX export — TDD against the produced bytes via python-docx round-trip.

The exported manuscript is a single .docx with:
  - title page (project title as Heading 1; study type + citation style),
  - one heading per section in canonical order,
  - prose paragraphs (basic HTML walker),
  - tables rendered from `<table>`,
  - bibliography section at the end with hanging-indent paragraphs.

PRISMA <img data:image/svg+xml> blocks are replaced with a placeholder
caption paragraph — the DOCX path skips SVG embedding because cairo isn't
installed in dev. The PDF export path keeps the SVG natively (reportlab
handles it).
"""
from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone

import pytest
from docx import Document

from research_api.db.models import Project
from research_api.services.export.bibliography import BibliographyEntry
from research_api.services.export.docx_export import (
    PRISMA_PLACEHOLDER,
    render_docx,
)


@dataclass
class Section:
    section_name: str
    content: str


@dataclass
class A:
    title: str | None = None
    authors: list[str] = field(default_factory=list)
    year: int | None = None
    journal: str | None = None
    doi: str | None = None
    volume: str | None = None
    issue: str | None = None
    pages: str | None = None


def _project(citation_style: str = "vancouver", title: str = "Test Project") -> Project:
    p = Project(
        id="proj1", user_id="user-a", title=title,
        study_type="Outcome Study", citation_style=citation_style,
        ai_provider="gemini",
    )
    p.created_at = datetime(2025, 1, 1, tzinfo=timezone.utc)
    p.updated_at = datetime(2025, 1, 1, tzinfo=timezone.utc)
    return p


def _parse(blob: bytes) -> Document:
    return Document(io.BytesIO(blob))


def test_render_docx_returns_valid_zip():
    blob = render_docx(project=_project(), sections=[], bibliography=[])
    assert isinstance(blob, bytes)
    assert blob[:2] == b"PK"  # ZIP magic
    with io.BytesIO(blob) as f:
        assert zipfile.is_zipfile(f)


def test_render_docx_round_trips_via_python_docx():
    blob = render_docx(project=_project(), sections=[], bibliography=[])
    doc = _parse(blob)
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("Test Project" in t for t in paragraphs)


def test_render_docx_title_page_carries_title_and_style():
    blob = render_docx(project=_project(citation_style="ieee"), sections=[], bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Project" in text
    assert "Outcome Study" in text
    # Citation style indicator present.
    assert "IEEE" in text.upper() or "ieee" in text


def test_render_docx_renders_sections_in_canonical_order():
    sections = [
        Section("Conclusion", "<p>Wrap up.</p>"),
        Section("Introduction", "<p>Intro prose.</p>"),
        Section("Results", "<p>Results prose.</p>"),
    ]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    # Section headings appear in canonical order.
    seen: list[str] = []
    for p in doc.paragraphs:
        for name in ("Introduction", "Methodology", "Results", "Discussion", "Conclusion", "Abstract"):
            if p.text.strip() == name:
                seen.append(name)
                break
    # Introduction must precede Results, which must precede Conclusion.
    assert seen.index("Introduction") < seen.index("Results") < seen.index("Conclusion")


def test_render_docx_empty_section_shows_placeholder():
    sections = [Section("Introduction", "")]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Introduction" in text
    assert "(Empty)" in text


def test_render_docx_renders_table_from_html():
    html = (
        "<table>"
        "<tr><th>Study</th><th>RoB</th></tr>"
        "<tr><td>Smith 2020</td><td>Low</td></tr>"
        "<tr><td>Jones 2021</td><td>High</td></tr>"
        "</table>"
    )
    sections = [Section("Results", html)]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    assert len(doc.tables) >= 1
    t = doc.tables[0]
    assert len(t.rows) == 3
    assert t.rows[0].cells[0].text.strip() == "Study"
    assert t.rows[1].cells[0].text.strip() == "Smith 2020"
    assert t.rows[2].cells[1].text.strip() == "High"


def test_render_docx_strips_script_tags():
    sections = [Section("Introduction", "<p>safe</p><script>alert('xss')</script>")]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "alert" not in text
    assert "safe" in text


def test_render_docx_renders_bibliography_entries():
    bib = [
        BibliographyEntry(article_id="a1", number=1, formatted="Doe J. A study. JAMA. 2024;1:1-2."),
        BibliographyEntry(article_id="a2", number=2, formatted="Smith K. Another. NEJM. 2023;2:3-4."),
    ]
    blob = render_docx(project=_project(), sections=[], bibliography=bib)
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "References" in text
    assert "1. Doe J" in text or "1.\tDoe J" in text or text.find("Doe J. A study") != -1
    assert "Smith K" in text


def test_render_docx_handles_empty_bibliography():
    blob = render_docx(project=_project(), sections=[], bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    # No "References" header for empty bib (or it's there as an empty stub —
    # accept either, but the file must parse cleanly).
    assert isinstance(text, str)


def test_render_docx_substitutes_inline_citation_sup():
    sections = [Section(
        "Introduction",
        '<p>Foo <sup data-citation data-article-id="a1">[1]</sup> bar.</p>',
    )]
    bib = [BibliographyEntry(article_id="a1", number=1, formatted="Doe J. ...")]
    blob = render_docx(project=_project(citation_style="ieee"),
                       sections=sections, bibliography=bib)
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    # The visible citation text must survive (in IEEE style it's `[1]`).
    assert "[1]" in text


def test_render_docx_replaces_svg_img_with_placeholder():
    # Embedded PRISMA SVG image is replaced with a placeholder paragraph in DOCX.
    svg_b64 = "PHN2Zy8+"  # `<svg/>`
    html = (
        "<p>Methods text.</p>"
        f'<img src="data:image/svg+xml;base64,{svg_b64}"/>'
    )
    sections = [Section("Methodology", html)]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert PRISMA_PLACEHOLDER in text


def test_render_docx_page_setup_a4_with_margins():
    blob = render_docx(project=_project(), sections=[], bibliography=[])
    doc = _parse(blob)
    section = doc.sections[0]
    # A4 is 210mm x 297mm — python-docx exposes these in EMU.
    # 210mm == 7.5594 inches; 297mm == 11.6929 inches.
    # We just sanity-check both dimensions are within tolerance of A4.
    width_in = section.page_width / 914400  # EMU per inch
    height_in = section.page_height / 914400
    assert 7.0 < width_in < 8.5
    assert 11.0 < height_in < 12.0
    # Margins ~ 1 inch.
    assert abs(section.top_margin / 914400 - 1.0) < 0.1
    assert abs(section.left_margin / 914400 - 1.0) < 0.1


def test_render_docx_handles_unknown_html_tags_gracefully():
    sections = [Section(
        "Introduction",
        "<p>Hello <unknown>world</unknown> <em>emphasis</em>.</p>",
    )]
    blob = render_docx(project=_project(), sections=sections, bibliography=[])
    doc = _parse(blob)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Hello" in text
    assert "world" in text  # unknown tag content surfaced
    assert "emphasis" in text
