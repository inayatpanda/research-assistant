"""Phase 20 (MP20) — Completed checklist export as PDF + DOCX.

Both functions are pure: they take a small payload and return bytes. The
route layer is responsible for HTTP headers and Content-Disposition.

Layout:
  * Title page with checklist name + run title + compliance summary
  * Table of items: ID | Item | Status | Section mapped | Comment
  * Footer line with the same compliance summary

Statuses are colour-coded in both formats:
  pass    → green
  fail    → red
  unclear → amber
  na      → grey
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


_STATUS_LABELS: dict[str, str] = {
    "pass": "Pass",
    "fail": "Fail",
    "unclear": "Unclear",
    "na": "N/A",
}

_STATUS_REPORTLAB: dict[str, colors.Color] = {
    "pass": colors.HexColor("#16a34a"),
    "fail": colors.HexColor("#dc2626"),
    "unclear": colors.HexColor("#d97706"),
    "na": colors.HexColor("#6b7280"),
}

_STATUS_DOCX: dict[str, RGBColor] = {
    "pass": RGBColor(0x16, 0xA3, 0x4A),
    "fail": RGBColor(0xDC, 0x26, 0x26),
    "unclear": RGBColor(0xD9, 0x77, 0x06),
    "na": RGBColor(0x6B, 0x72, 0x80),
}


def _summary_counts(items: list[dict[str, Any]]) -> tuple[int, int, int, int]:
    """Return (n_pass, n_fail, n_unclear, n_na)."""
    n_pass = sum(1 for i in items if (i or {}).get("status") == "pass")
    n_fail = sum(1 for i in items if (i or {}).get("status") == "fail")
    n_unclear = sum(1 for i in items if (i or {}).get("status") == "unclear")
    n_na = sum(1 for i in items if (i or {}).get("status") == "na")
    return n_pass, n_fail, n_unclear, n_na


def _summary_line(pct: float, items: list[dict[str, Any]]) -> str:
    n_pass, n_fail, n_unclear, n_na = _summary_counts(items)
    return (
        f"{pct:.1f}% compliance "
        f"({n_pass} passed, {n_fail} failed, {n_unclear} unclear, {n_na} N/A)"
    )


def render_pdf(
    *,
    checklist_name: str,
    run_title: str,
    items: list[dict[str, Any]],
    compliance_pct: float,
) -> bytes:
    """Render the run to a single-document PDF."""
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "TitleX",
            parent=base["Title"],
            fontSize=18,
            spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontSize=12, spaceAfter=6
        ),
        "body": ParagraphStyle(
            "Body", parent=base["BodyText"], fontSize=8, alignment=TA_LEFT,
            leading=11,
        ),
        "cell": ParagraphStyle(
            "Cell", parent=base["BodyText"], fontSize=8, alignment=TA_LEFT,
            leading=10, spaceBefore=0, spaceAfter=0,
        ),
    }

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title=f"{checklist_name} — {run_title}",
    )

    summary = _summary_line(compliance_pct, items)
    story: list[Any] = [
        Paragraph(_html_escape(checklist_name), styles["title"]),
        Paragraph(_html_escape(run_title), styles["h2"]),
        Paragraph(_html_escape(summary), styles["body"]),
        Spacer(1, 8),
    ]

    # Header row.
    header = ["ID", "Item", "Status", "Section", "Comment"]
    rows: list[list[Any]] = [
        [Paragraph(f"<b>{h}</b>", styles["cell"]) for h in header]
    ]
    status_colors: list[tuple[int, int, colors.Color]] = []
    for row_idx, raw in enumerate(items, start=1):
        item = raw or {}
        status_raw = str(item.get("status") or "unclear")
        status_text = _STATUS_LABELS.get(status_raw, status_raw)
        cells = [
            Paragraph(_html_escape(str(item.get("item_id") or "")), styles["cell"]),
            Paragraph(_html_escape(str(item.get("item_text") or "")), styles["cell"]),
            Paragraph(f"<b>{_html_escape(status_text)}</b>", styles["cell"]),
            Paragraph(
                _html_escape(str(item.get("mapped_section") or "")),
                styles["cell"],
            ),
            Paragraph(
                _html_escape(str(item.get("comment") or "")),
                styles["cell"],
            ),
        ]
        rows.append(cells)
        status_colors.append(
            (row_idx, 2, _STATUS_REPORTLAB.get(status_raw, colors.black))
        )

    table = Table(
        rows,
        colWidths=[
            0.6 * inch,
            2.6 * inch,
            0.7 * inch,
            1.0 * inch,
            2.6 * inch,
        ],
        repeatRows=1,
    )
    style_cmds: list[Any] = [
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2ff")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ]
    for r, c, color in status_colors:
        style_cmds.append(("TEXTCOLOR", (c, r), (c, r), color))
    table.setStyle(TableStyle(style_cmds))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph(_html_escape(summary), styles["body"]))

    doc.build(story)
    return buf.getvalue()


def render_docx(
    *,
    checklist_name: str,
    run_title: str,
    items: list[dict[str, Any]],
    compliance_pct: float,
) -> bytes:
    """Render the run to a DOCX byte stream."""
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run(checklist_name)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    sub = subtitle.add_run(run_title)
    sub.italic = True
    sub.font.size = Pt(12)

    summary = _summary_line(compliance_pct, items)
    doc.add_paragraph(summary)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # pragma: no cover - style availability varies
        pass
    hdr = table.rows[0].cells
    for idx, label in enumerate(["ID", "Item", "Status", "Section", "Comment"]):
        p = hdr[idx].paragraphs[0]
        p.add_run(label).bold = True

    for raw in items:
        item = raw or {}
        status_raw = str(item.get("status") or "unclear")
        row = table.add_row().cells
        row[0].text = str(item.get("item_id") or "")
        row[1].text = str(item.get("item_text") or "")
        status_para = row[2].paragraphs[0]
        status_run = status_para.add_run(
            _STATUS_LABELS.get(status_raw, status_raw)
        )
        status_run.bold = True
        status_run.font.color.rgb = _STATUS_DOCX.get(status_raw, RGBColor(0, 0, 0))
        row[3].text = str(item.get("mapped_section") or "")
        row[4].text = str(item.get("comment") or "")

    doc.add_paragraph(summary)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _html_escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


__all__ = ["render_docx", "render_pdf"]
