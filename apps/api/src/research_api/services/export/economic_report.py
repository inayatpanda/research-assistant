"""Phase 18 (MP18) — CHEERS 2022 economic evaluation report builder.

Builds DOCX / PDF for a single ``EconomicAnalysis`` + its ``EconomicResult``,
laid out against the CHEERS 2022 checklist (Husereau et al. Value Health
2022):

  * Title page
  * Background & objectives
  * Methods — population, time horizon, perspective, discount rates, value
    set, analytic methods (bootstrap), software
  * Results — mean cost diff, mean QALY diff, ICER, dominance, plane
    figure, CEAC figure
  * Sensitivity — PSA / DSA / scenario summaries
  * Discussion / conclusion

The two PNGs (plane + CEAC) are embedded directly from their data URIs.
"""
from __future__ import annotations

import base64
import io
from dataclasses import dataclass
from typing import Any, Protocol

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


class _ProjectLike(Protocol):
    title: str
    study_type: str


@dataclass
class CHEERSContext:
    name: str
    perspective: str
    time_horizon_months: int
    currency: str
    discount_rate_costs: float
    discount_rate_qalys: float
    intervention_label: str
    comparator_label: str
    value_set: str
    bootstrap_n: int
    seed: int
    mean_cost_diff: float
    mean_qaly_diff: float
    icer: float | None
    dominance_status: str
    nmb_at_thresholds: dict[str, Any]
    wtp_thresholds: list[int]
    ai_interpretation: str | None
    plane_png_uri: str
    ceac_png_uri: str
    sensitivity: dict[str, Any] | None


def _decode_data_uri(uri: str) -> bytes:
    """Return the raw PNG bytes from a ``data:image/png;base64,...`` URI."""
    if not uri or "," not in uri:
        return b""
    _, b64 = uri.split(",", 1)
    try:
        return base64.b64decode(b64)
    except Exception:
        return b""


def _fmt_icer(icer: float | None, currency: str) -> str:
    if icer is None:
        return "n/a (dominance applies)"
    return f"{currency} {icer:,.0f} / QALY"


def _docx_bytes(project: _ProjectLike, ctx: CHEERSContext) -> bytes:
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Economic Evaluation Report (CHEERS 2022)")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(project.title or "(untitled project)").italic = True
    doc.add_paragraph(f"Analysis: {ctx.name}")

    # Methods
    h = doc.add_paragraph()
    h.add_run("Methods").bold = True
    doc.add_paragraph(f"Perspective: {ctx.perspective}.")
    doc.add_paragraph(
        f"Time horizon: {ctx.time_horizon_months} months. "
        f"Currency: {ctx.currency}."
    )
    doc.add_paragraph(
        f"Discount rate (costs): {ctx.discount_rate_costs:.3f}. "
        f"Discount rate (QALYs): {ctx.discount_rate_qalys:.3f}."
    )
    doc.add_paragraph(f"Utility value set: {ctx.value_set}.")
    doc.add_paragraph(
        f"Bootstrap replicates: {ctx.bootstrap_n} (seed {ctx.seed})."
    )
    doc.add_paragraph(
        f"Comparison: {ctx.intervention_label} vs {ctx.comparator_label}."
    )

    # Results
    h = doc.add_paragraph()
    h.add_run("Results").bold = True
    doc.add_paragraph(
        f"Mean incremental cost: {ctx.currency} {ctx.mean_cost_diff:,.2f}"
    )
    doc.add_paragraph(
        f"Mean incremental QALYs: {ctx.mean_qaly_diff:.4f}"
    )
    doc.add_paragraph(
        f"ICER: {_fmt_icer(ctx.icer, ctx.currency)}"
    )
    doc.add_paragraph(f"Dominance: {ctx.dominance_status}")
    if ctx.nmb_at_thresholds:
        items = sorted(ctx.nmb_at_thresholds.items(), key=lambda kv: int(kv[0]))
        for k, v in items:
            doc.add_paragraph(
                f"  NMB at {ctx.currency} {int(k):,}/QALY: {float(v):,.0f}"
            )

    # Embedded plane PNG
    plane_bytes = _decode_data_uri(ctx.plane_png_uri)
    if plane_bytes:
        doc.add_paragraph().add_run("Cost-effectiveness plane").bold = True
        doc.add_picture(io.BytesIO(plane_bytes), width=Inches(5.5))

    # Embedded CEAC PNG
    ceac_bytes = _decode_data_uri(ctx.ceac_png_uri)
    if ceac_bytes:
        doc.add_paragraph().add_run(
            "Cost-effectiveness acceptability curve"
        ).bold = True
        doc.add_picture(io.BytesIO(ceac_bytes), width=Inches(5.5))

    # Sensitivity
    if ctx.sensitivity:
        h = doc.add_paragraph()
        h.add_run("Sensitivity analysis").bold = True
        kind = ctx.sensitivity.get("type", "(unspecified)")
        doc.add_paragraph(f"Kind: {kind}")
        summary = ctx.sensitivity.get("summary") or {}
        for k, v in summary.items():
            doc.add_paragraph(f"  {k}: {v}")

    # Discussion / AI prose
    if ctx.ai_interpretation:
        h = doc.add_paragraph()
        h.add_run("Discussion").bold = True
        doc.add_paragraph(ctx.ai_interpretation)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes(project: _ProjectLike, ctx: CHEERSContext) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=0.75 * inch, rightMargin=0.75 * inch
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CHEERSTitle",
        parent=styles["Title"],
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "CHEERSSubtitle",
        parent=styles["BodyText"],
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    h_style = ParagraphStyle(
        "CHEERSSection",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = styles["BodyText"]
    elements: list[Any] = [
        Paragraph("Economic Evaluation Report (CHEERS 2022)", title_style),
        Paragraph(project.title or "(untitled project)", subtitle_style),
        Paragraph(f"Analysis: {ctx.name}", body_style),
        Spacer(1, 0.15 * inch),
        Paragraph("Methods", h_style),
    ]
    rows = [
        ["Perspective", ctx.perspective],
        ["Time horizon", f"{ctx.time_horizon_months} months"],
        ["Currency", ctx.currency],
        ["Discount rate (costs)", f"{ctx.discount_rate_costs:.3f}"],
        ["Discount rate (QALYs)", f"{ctx.discount_rate_qalys:.3f}"],
        ["Utility value set", ctx.value_set],
        ["Bootstrap", f"{ctx.bootstrap_n} reps (seed {ctx.seed})"],
        [
            "Comparison",
            f"{ctx.intervention_label} vs {ctx.comparator_label}",
        ],
    ]
    common_table_style = TableStyle(
        [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONT", (0, 0), (0, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]
    )
    table = Table(rows, colWidths=[1.7 * inch, 4.3 * inch])
    table.setStyle(common_table_style)
    elements.append(table)
    elements.append(Spacer(1, 0.15 * inch))

    elements.append(Paragraph("Results", h_style))
    res_rows = [
        [
            "Mean incremental cost",
            f"{ctx.currency} {ctx.mean_cost_diff:,.2f}",
        ],
        ["Mean incremental QALYs", f"{ctx.mean_qaly_diff:.4f}"],
        ["ICER", _fmt_icer(ctx.icer, ctx.currency)],
        ["Dominance", ctx.dominance_status],
    ]
    if ctx.nmb_at_thresholds:
        items = sorted(
            ctx.nmb_at_thresholds.items(), key=lambda kv: int(kv[0])
        )
        for k, v in items:
            res_rows.append(
                [
                    f"NMB at {ctx.currency} {int(k):,}/QALY",
                    f"{float(v):,.0f}",
                ]
            )
    res_table = Table(res_rows, colWidths=[1.7 * inch, 4.3 * inch])
    res_table.setStyle(common_table_style)
    elements.append(res_table)
    elements.append(Spacer(1, 0.15 * inch))

    plane_bytes = _decode_data_uri(ctx.plane_png_uri)
    if plane_bytes:
        elements.append(Paragraph("Cost-effectiveness plane", h_style))
        elements.append(Image(io.BytesIO(plane_bytes), width=5.0 * inch, height=3.6 * inch))
        elements.append(Spacer(1, 0.1 * inch))
    ceac_bytes = _decode_data_uri(ctx.ceac_png_uri)
    if ceac_bytes:
        elements.append(
            Paragraph("Cost-effectiveness acceptability curve", h_style)
        )
        elements.append(Image(io.BytesIO(ceac_bytes), width=5.0 * inch, height=3.6 * inch))
        elements.append(Spacer(1, 0.1 * inch))

    if ctx.sensitivity:
        elements.append(Paragraph("Sensitivity analysis", h_style))
        kind = ctx.sensitivity.get("type", "(unspecified)")
        elements.append(Paragraph(f"Kind: {kind}", body_style))
        summary = ctx.sensitivity.get("summary") or {}
        for k, v in summary.items():
            elements.append(Paragraph(f"{k}: {v}", body_style))

    if ctx.ai_interpretation:
        elements.append(Paragraph("Discussion", h_style))
        elements.append(Paragraph(ctx.ai_interpretation, body_style))

    doc.build(elements)
    return buf.getvalue()


def build_economic_report(
    project: _ProjectLike, ctx: CHEERSContext, *, fmt: str = "docx"
) -> bytes:
    if fmt not in ("docx", "pdf"):
        raise ValueError("fmt must be 'docx' or 'pdf'")
    if fmt == "docx":
        return _docx_bytes(project, ctx)
    return _pdf_bytes(project, ctx)


__all__ = ["build_economic_report", "CHEERSContext"]
