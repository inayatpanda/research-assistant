"""DOCX manuscript export via python-docx.

Layout:
  - title page (project title as H1; study type + citation style).
  - Phase 10 ICMJE block (when supplied): authors with affiliation
    superscripts, corresponding-author email, affiliation list, COI/funding/
    ethics statements, optional structured abstract (Background/Methods/
    Results/Conclusions) replacing the freeform Abstract section.
  - one heading per section in canonical order (Abstract → Conclusion).
  - prose paragraphs (HTML parsed via `_html_walker`).
  - tables from `<table>` blocks.
  - References section at the end with hanging-indent paragraphs.

PRISMA `<img data:image/svg+xml>` images are replaced with a placeholder
caption paragraph — the DOCX export skips raster conversion because the dev
machine ships without cairo. The PDF export embeds the same SVGs natively.
"""
from __future__ import annotations

import html
import io
from dataclasses import dataclass, field
from typing import Iterable, Protocol

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from ._html_walker import walk_html
from .bibliography import CANONICAL_SECTION_ORDER, BibliographyEntry


SUPERSCRIPTS = "⁰¹²³⁴⁵⁶⁷⁸⁹"


def _to_superscript(n: int) -> str:
    """Render a positive integer as Unicode superscript digits."""
    return "".join(SUPERSCRIPTS[int(d)] for d in str(n))


@dataclass
class FrontMatterPayload:
    """Phase 10 ICMJE block injected before the manuscript body.

    Pure data — no ORM references — so it round-trips through the export
    routes without dragging SQLAlchemy in.
    """
    authors: list[dict] = field(default_factory=list)
    """Each dict: {full_name, given_name, family_name, orcid, email,
    is_corresponding, position, affiliation_ids}."""
    affiliations: list[dict] = field(default_factory=list)
    """Each dict: {id, name, address, city, country, position}."""
    funding_statement: str | None = None
    funders: list[dict] = field(default_factory=list)
    """Each dict: {name, grant_id}."""
    ethics_irb: str | None = None
    ethics_approval_number: str | None = None
    ethics_consent: str | None = None
    conflicts_statement: str | None = None
    structured_abstract_enabled: bool = False
    structured_abstract: dict | None = None
    """Dict: {background, methods, results, conclusions}."""

PRISMA_PLACEHOLDER = (
    "Figure: PRISMA 2020 flow diagram (rendered in the online manuscript)."
)


class _ProjectLike(Protocol):
    title: str
    study_type: str
    citation_style: str


class _SectionLike(Protocol):
    section_name: str
    content: str


def _style_label(style: str) -> str:
    return {"vancouver": "Vancouver", "apa": "APA 7th",
            "harvard": "Harvard", "ieee": "IEEE"}.get(style, style)


def _configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.27)   # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def _add_title_page(
    document: Document,
    project: _ProjectLike,
    frontmatter: FrontMatterPayload | None = None,
) -> None:
    title = document.add_heading(project.title, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if frontmatter and frontmatter.authors:
        _add_authors_block(document, frontmatter)
    else:
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(f"Study type: {project.study_type}").italic = True
        p2 = document.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p2.add_run(f"Citation style: {_style_label(project.citation_style)}").italic = True
    document.add_paragraph()  # blank spacer


def _build_affiliation_numbering(
    fm: FrontMatterPayload,
) -> tuple[dict[str, int], list[dict]]:
    """Walk authors in position order; assign 1-based numbers to each
    affiliation in first-encounter order.

    Returns:
        (aff_id → number map, ordered affiliation rows)
    """
    aff_by_id = {a["id"]: a for a in fm.affiliations}
    seen: dict[str, int] = {}
    ordered: list[dict] = []
    for author in sorted(fm.authors, key=lambda a: a.get("position", 0)):
        for aff_id in author.get("affiliation_ids") or []:
            if aff_id in seen or aff_id not in aff_by_id:
                continue
            seen[aff_id] = len(seen) + 1
            ordered.append(aff_by_id[aff_id])
    return seen, ordered


def _add_authors_block(document: Document, fm: FrontMatterPayload) -> None:
    """Authors line + numbered affiliations + corresponding-author line.

    All user-supplied strings flow through `html.escape` for defence in depth
    even though python-docx doesn't render markup — this keeps the same
    sanitisation discipline whether output lands in DOCX or PDF.
    """
    aff_numbering, ordered_affs = _build_affiliation_numbering(fm)
    authors_sorted = sorted(fm.authors, key=lambda a: a.get("position", 0))
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for idx, author in enumerate(authors_sorted):
        if idx > 0:
            p.add_run(", ")
        name = html.escape(author.get("full_name") or "")
        p.add_run(name)
        nums = [
            aff_numbering[aid]
            for aid in (author.get("affiliation_ids") or [])
            if aid in aff_numbering
        ]
        if nums:
            sup_text = ",".join(str(n) for n in sorted(nums))
            sup_run = p.add_run(sup_text)
            sup_run.font.superscript = True
        if author.get("is_corresponding"):
            star_run = p.add_run("*")
            star_run.font.superscript = True

    if ordered_affs:
        for aff in ordered_affs:
            ap = document.add_paragraph()
            number = aff_numbering[aff["id"]]
            ap.add_run(_to_superscript(number) + " ")
            parts = [
                html.escape(aff.get("name") or ""),
                html.escape(aff.get("address") or ""),
                html.escape(aff.get("city") or ""),
                html.escape(aff.get("country") or ""),
            ]
            text = ", ".join(p for p in parts if p)
            ap.add_run(text).italic = True

    corresponding = next(
        (a for a in authors_sorted if a.get("is_corresponding")), None
    )
    if corresponding and corresponding.get("email"):
        cp = document.add_paragraph()
        cp.add_run("* Corresponding author: ").bold = True
        cp.add_run(html.escape(corresponding["email"] or ""))


def _add_frontmatter_statements(
    document: Document, fm: FrontMatterPayload
) -> None:
    """Emit COI / funding / ethics paragraphs after the body's References."""
    blocks: list[tuple[str, str | None]] = [
        ("Conflicts of Interest", fm.conflicts_statement),
        ("Funding", _compose_funding(fm)),
        ("Ethics", _compose_ethics(fm)),
    ]
    emitted = False
    for label, text in blocks:
        if not text:
            continue
        if not emitted:
            document.add_paragraph()
            emitted = True
        h = document.add_paragraph()
        h.add_run(label).bold = True
        body = document.add_paragraph()
        body.add_run(html.escape(text))


def _compose_funding(fm: FrontMatterPayload) -> str | None:
    parts: list[str] = []
    if fm.funding_statement:
        parts.append(fm.funding_statement.strip())
    if fm.funders:
        chunks = []
        for f in fm.funders:
            name = (f.get("name") or "").strip()
            grant = (f.get("grant_id") or "").strip()
            if not name:
                continue
            chunks.append(f"{name} ({grant})" if grant else name)
        if chunks:
            parts.append("Funders: " + "; ".join(chunks) + ".")
    return " ".join(parts) or None


def _compose_ethics(fm: FrontMatterPayload) -> str | None:
    parts: list[str] = []
    if fm.ethics_irb:
        parts.append(f"IRB: {fm.ethics_irb.strip()}.")
    if fm.ethics_approval_number:
        parts.append(
            f"Approval number: {fm.ethics_approval_number.strip()}."
        )
    if fm.ethics_consent:
        parts.append(fm.ethics_consent.strip())
    return " ".join(parts) or None


def _add_structured_abstract(
    document: Document, abstract: dict
) -> None:
    """Replace the freeform Abstract section with 4 labelled paragraphs."""
    document.add_heading("Abstract", level=1)
    labels = [
        ("Background", abstract.get("background", "") or ""),
        ("Methods", abstract.get("methods", "") or ""),
        ("Results", abstract.get("results", "") or ""),
        ("Conclusions", abstract.get("conclusions", "") or ""),
    ]
    for label, content in labels:
        para = document.add_paragraph()
        para.add_run(label + ": ").bold = True
        para.add_run(html.escape(content) or "(Empty)")


def _order_sections(sections: Iterable[_SectionLike]) -> list[_SectionLike]:
    by_name = {s.section_name: s for s in sections}
    return [by_name[n] for n in CANONICAL_SECTION_ORDER if n in by_name]


def _add_run(paragraph, text: str, styles: set[str]) -> None:
    run = paragraph.add_run(text)
    if "bold" in styles:
        run.bold = True
    if "italic" in styles:
        run.italic = True
    if "sup" in styles:
        run.font.superscript = True
    if "underline" in styles:
        run.underline = True


def _render_events(document: Document, events: list[tuple]) -> None:
    """Walk the event stream and emit DOCX paragraphs / tables."""
    paragraph = None
    table = None
    row_cells: list = []
    cell_idx = 0
    pending_cell_paragraph = None

    def flush_paragraph() -> None:
        nonlocal paragraph
        paragraph = None

    for ev in events:
        kind = ev[0]
        if kind == "paragraph_start":
            paragraph = document.add_paragraph()
        elif kind == "paragraph_end":
            flush_paragraph()
        elif kind == "heading_start":
            level = ev[1]
            paragraph = document.add_heading("", level=min(max(level, 1), 4))
        elif kind == "heading_end":
            flush_paragraph()
        elif kind == "text":
            _, text, styles = ev
            if pending_cell_paragraph is not None:
                _add_run(pending_cell_paragraph, text, styles)
            elif paragraph is not None:
                _add_run(paragraph, text, styles)
        elif kind == "svg_img":
            flush_paragraph()
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(PRISMA_PLACEHOLDER)
            run.italic = True
        elif kind == "table_start":
            flush_paragraph()
            table = document.add_table(rows=0, cols=0)
            table.style = "Table Grid"
        elif kind == "row_start":
            if table is None:
                continue
            row = table.add_row()
            row_cells = list(row.cells)
            cell_idx = 0
        elif kind == "row_end":
            row_cells = []
            cell_idx = 0
        elif kind == "cell_start":
            is_header = ev[1]
            if table is None:
                continue
            # Add a column on the fly if the row needs more cells than exist.
            while cell_idx >= len(row_cells):
                # python-docx requires explicit column add to extend the row.
                if not table.columns or cell_idx >= len(table.columns):
                    table.add_column(Inches(1.0))
                row_cells = list(table.rows[-1].cells)
            cell = row_cells[cell_idx]
            cell.text = ""  # clear default empty paragraph
            pending_cell_paragraph = cell.paragraphs[0]
            if is_header:
                for run in pending_cell_paragraph.runs:
                    run.bold = True
        elif kind == "cell_end":
            pending_cell_paragraph = None
            cell_idx += 1
        elif kind == "table_end":
            table = None


def _add_section(document: Document, section: _SectionLike) -> None:
    document.add_heading(section.section_name, level=1)
    events = walk_html(section.content or "")
    if not events:
        p = document.add_paragraph()
        p.add_run("(Empty)").italic = True
        return
    _render_events(document, events)


def _add_bibliography(document: Document, entries: list[BibliographyEntry]) -> None:
    if not entries:
        return
    document.add_page_break()
    document.add_heading("References", level=1)
    for entry in entries:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(f"{entry.number}. {entry.formatted}")


def html_to_docx_bytes(html_str: str, *, title: str | None = None) -> bytes:
    """Render a standalone HTML blob (e.g. cover-letter body) to DOCX bytes.

    Used by the submission-package builder for cover_letter.docx and
    response_to_reviewers.docx. The body is walked through `_html_walker`
    so the same `<p>`/`<strong>`/`<em>`/`<table>` permissive grammar
    applies, giving identical rendering to the main manuscript export.
    """
    document = Document()
    _configure_page(document)
    if title:
        h = document.add_heading(title, level=0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    events = walk_html(html_str or "")
    if not events:
        p = document.add_paragraph()
        p.add_run("(Empty)").italic = True
    else:
        _render_events(document, events)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def tables_to_individual_docx(section_html: str) -> dict[int, bytes]:
    """Phase 12 — extract every `<table>` from a section's HTML and render
    each into its own standalone DOCX, keyed by the table's 1-based index
    within the section.

    The submission package writer concatenates results across all sections
    and renumbers tables globally so the filenames stay `Table_1.docx`,
    `Table_2.docx`, ... in document order.
    """
    events = list(walk_html(section_html or ""))
    out: dict[int, bytes] = {}
    # Split the event stream into per-table sub-streams.
    table_idx = 0
    current: list[tuple] | None = None
    for ev in events:
        kind = ev[0]
        if kind == "table_start":
            table_idx += 1
            current = [ev]
        elif current is not None:
            current.append(ev)
            if kind == "table_end":
                doc = Document()
                _configure_page(doc)
                _render_events(doc, current)
                buf = io.BytesIO()
                doc.save(buf)
                out[table_idx] = buf.getvalue()
                current = None
    return out


def render_docx(
    *,
    project: _ProjectLike,
    sections: Iterable[_SectionLike],
    bibliography: list[BibliographyEntry],
    frontmatter: FrontMatterPayload | None = None,
) -> bytes:
    """Return the DOCX bytes for the manuscript.

    When `frontmatter` is supplied (Phase 10 ICMJE block), the title page is
    replaced with the structured authors + affiliations layout, the freeform
    Abstract section is swapped for a structured-abstract block if the
    project has opted in, and trailing COI/Funding/Ethics paragraphs are
    appended after References.
    """
    document = Document()
    _configure_page(document)
    _add_title_page(document, project, frontmatter)

    use_structured = bool(
        frontmatter
        and frontmatter.structured_abstract_enabled
        and frontmatter.structured_abstract
    )
    for section in _order_sections(sections):
        if use_structured and section.section_name == "Abstract":
            _add_structured_abstract(document, frontmatter.structured_abstract)  # type: ignore[union-attr]
        else:
            _add_section(document, section)
    _add_bibliography(document, bibliography)
    if frontmatter:
        _add_frontmatter_statements(document, frontmatter)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
