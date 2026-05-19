"""Phase 17 (MP17) — Statistical Analysis Plan (SAP) export.

Two helpers:

  * ``compute_integrity_hash(steps)``: deterministic SHA-256 of the plan's
    steps list. The encoding choices:
      - JSON-dumped with ``sort_keys=True`` so dict-key insertion order is
        irrelevant
      - ``ensure_ascii=False`` (so identical Unicode rendering across
        platforms)
      - floats rounded to 8 decimal places before encoding (handles the
        usual numeric drift across NumPy/Python casts)
      - lists preserved in order (lists carry meaning in plan steps)
    This is the single source of truth for pre-registration locking.
  * ``build_sap_document(project, plan, *, fmt)``: returns DOCX / PDF bytes.

Plan step shapes mirror MP13.5's PlanStep — ``{type, args}`` — so the
renderer can iterate them generically.
"""
from __future__ import annotations

import hashlib
import io
import json
import math
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Iterable, Protocol

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


# ─── Integrity hash ────────────────────────────────────────────────────────


def _round_floats(value: Any, *, ndigits: int = 8) -> Any:
    """Recursively round all floats inside a JSON-shaped structure.

    Keeps integers and ``bool`` untouched (Python ``bool`` is a subclass of
    int — must check it first to avoid casting to float). NaN / Inf are
    serialised as the strings ``"NaN"`` / ``"Infinity"`` / ``"-Infinity"``
    so different platforms produce the same bytes.
    """
    if isinstance(value, bool):
        return value
    if isinstance(value, float):
        if math.isnan(value):
            return "NaN"
        if math.isinf(value):
            return "Infinity" if value > 0 else "-Infinity"
        return round(float(value), ndigits)
    if isinstance(value, dict):
        return {k: _round_floats(v, ndigits=ndigits) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [_round_floats(v, ndigits=ndigits) for v in value]
    return value


def compute_integrity_hash(steps: list[dict[str, Any]]) -> str:
    """SHA-256 of the canonical-JSON-encoded steps list.

    Deterministic-encoding choices (locked in for the lifetime of the
    integrity hash):

      * ``sort_keys=True``
      * ``ensure_ascii=False``
      * ``separators=(",", ":")``
      * floats rounded to 8 decimal places
      * NaN / Inf encoded as fixed strings
    """
    rounded = _round_floats(steps)
    payload = json.dumps(
        rounded,
        sort_keys=True,
        ensure_ascii=False,
        separators=(",", ":"),
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


# ─── SAP document builder ──────────────────────────────────────────────────


class _ProjectLike(Protocol):
    title: str
    study_type: str


@dataclass
class SAPSection:
    """A single rendered analysis step for the SAP document."""

    index: int
    step_type: str
    title: str
    hypothesis: str
    primary_or_secondary: str
    variables: str
    population: str
    multiplicity_adjustment: str


def _step_label(step: dict[str, Any]) -> str:
    args = step.get("args") or {}
    kind = step.get("type", "?")
    if kind == "test":
        return str(args.get("test_key") or args.get("question_type") or "test")
    if kind == "transform":
        return str(args.get("op_type") or "transform")
    if kind == "plot":
        return str(args.get("geom") or "plot")
    return kind


def _sap_sections(steps: list[dict[str, Any]]) -> list[SAPSection]:
    out: list[SAPSection] = []
    for idx, step in enumerate(steps, start=1):
        args = step.get("args") or {}
        # Pull display-only fields with safe fallbacks.
        variables = args.get("variables") or {}
        if isinstance(variables, dict):
            var_str = ", ".join(
                f"{k}={v}" for k, v in variables.items() if v not in (None, "")
            ) or "(none)"
        else:
            var_str = str(variables)
        out.append(
            SAPSection(
                index=idx,
                step_type=step.get("type", "?"),
                title=_step_label(step),
                hypothesis=str(args.get("hypothesis") or "(not pre-specified)"),
                primary_or_secondary=str(args.get("primary_or_secondary") or "secondary"),
                variables=var_str,
                population=str(args.get("population") or "(whole dataset)"),
                multiplicity_adjustment=str(
                    args.get("multiplicity_adjustment") or "(none specified)"
                ),
            )
        )
    return out


def _docx_bytes(
    project: _ProjectLike,
    plan_name: str,
    integrity_hash: str,
    locked_at: datetime | None,
    sections: list[SAPSection],
) -> bytes:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Statistical Analysis Plan")
    run.bold = True
    run.font.size = Pt(22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(project.title or "(untitled project)").italic = True
    doc.add_paragraph(f"Plan: {plan_name}")
    doc.add_paragraph(f"Integrity hash: {integrity_hash}")
    if locked_at is not None:
        doc.add_paragraph(f"Locked at: {locked_at.isoformat()}")
    doc.add_paragraph(f"Total analyses: {len(sections)}")

    for sec in sections:
        h = doc.add_paragraph()
        h_run = h.add_run(f"{sec.index}. {sec.title} ({sec.step_type})")
        h_run.bold = True
        h_run.font.size = Pt(14)
        doc.add_paragraph(f"Hypothesis: {sec.hypothesis}")
        doc.add_paragraph(f"Primary / secondary: {sec.primary_or_secondary}")
        doc.add_paragraph(f"Variables: {sec.variables}")
        doc.add_paragraph(f"Population: {sec.population}")
        doc.add_paragraph(
            f"Multiplicity adjustment: {sec.multiplicity_adjustment}"
        )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes(
    project: _ProjectLike,
    plan_name: str,
    integrity_hash: str,
    locked_at: datetime | None,
    sections: list[SAPSection],
) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=0.75 * inch, rightMargin=0.75 * inch
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SAPTitle",
        parent=styles["Title"],
        fontSize=22,
        alignment=TA_CENTER,
        spaceAfter=14,
    )
    subtitle_style = ParagraphStyle(
        "SAPSubtitle",
        parent=styles["BodyText"],
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    h_style = ParagraphStyle(
        "SAPSection",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = styles["BodyText"]
    elements: list[Any] = [
        Paragraph("Statistical Analysis Plan", title_style),
        Paragraph(project.title or "(untitled project)", subtitle_style),
        Paragraph(f"Plan: {plan_name}", body_style),
        Paragraph(f"Integrity hash: <font face='Courier'>{integrity_hash}</font>", body_style),
    ]
    if locked_at is not None:
        elements.append(Paragraph(f"Locked at: {locked_at.isoformat()}", body_style))
    elements.append(Paragraph(f"Total analyses: {len(sections)}", body_style))
    elements.append(Spacer(1, 0.2 * inch))

    for sec in sections:
        elements.append(
            Paragraph(f"{sec.index}. {sec.title} ({sec.step_type})", h_style)
        )
        rows = [
            ["Hypothesis", sec.hypothesis],
            ["Primary / secondary", sec.primary_or_secondary],
            ["Variables", sec.variables],
            ["Population", sec.population],
            ["Multiplicity adjustment", sec.multiplicity_adjustment],
        ]
        table = Table(rows, colWidths=[1.6 * inch, 4.4 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONT", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        elements.append(table)
        elements.append(Spacer(1, 0.1 * inch))
    doc.build(elements)
    return buf.getvalue()


def build_sap_document(
    project: _ProjectLike,
    plan: Any,
    *,
    fmt: str = "docx",
) -> bytes:
    """Build a SAP document for ``plan`` of the given format.

    ``plan`` must expose ``name``, ``steps``, ``integrity_hash``, and
    ``locked_at`` (any of which may be falsy — we render defensively).
    """
    if fmt not in ("docx", "pdf"):
        raise ValueError("fmt must be 'docx' or 'pdf'")
    steps = list(plan.steps or [])
    sections = _sap_sections(steps)
    integrity = plan.integrity_hash or compute_integrity_hash(steps)
    locked_at: datetime | None = plan.locked_at
    if fmt == "docx":
        return _docx_bytes(project, plan.name, integrity, locked_at, sections)
    return _pdf_bytes(project, plan.name, integrity, locked_at, sections)


__all__ = ["compute_integrity_hash", "build_sap_document"]
