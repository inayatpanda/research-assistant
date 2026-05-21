"""Phase 4.6 — Render an AI peer-review critique to DOCX / PDF.

Both functions are pure: ``(critique_dict, source_title) -> bytes``. The
route layer handles HTTP headers + Content-Disposition.

Sections rendered (in order):

* Title block (source title + recommendation badge)
* Overall impression
* Strengths
* Major issues
* Minor issues
* Methodological / statistical / reporting / presentation / references concerns
* Suggestions for improvement
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

_REC_LABELS: dict[str, str] = {
    "reject": "Reject",
    "major_revision": "Major Revision",
    "minor_revision": "Minor Revision",
    "accept": "Accept",
}

_REC_COLORS_HEX: dict[str, str] = {
    "reject": "#dc2626",
    "major_revision": "#d97706",
    "minor_revision": "#2563eb",
    "accept": "#16a34a",
}

_REC_COLORS_DOCX: dict[str, RGBColor] = {
    "reject": RGBColor(0xDC, 0x26, 0x26),
    "major_revision": RGBColor(0xD9, 0x77, 0x06),
    "minor_revision": RGBColor(0x25, 0x63, 0xEB),
    "accept": RGBColor(0x16, 0xA3, 0x4A),
}

# Order of structured-output sections, with display labels.
_SECTIONS: tuple[tuple[str, str], ...] = (
    ("strengths", "Strengths"),
    ("major_issues", "Major Issues"),
    ("minor_issues", "Minor Issues"),
    ("methodological_concerns", "Methodological Concerns"),
    ("statistical_concerns", "Statistical Concerns"),
    ("reporting_concerns", "Reporting Concerns"),
    ("presentation_concerns", "Presentation Concerns"),
    ("references_concerns", "References Concerns"),
    ("suggestions_for_improvement", "Suggestions for Improvement"),
)


def _html_escape(value: str) -> str:
    return (
        (value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def render_critique_pdf(*, source_title: str, critique: dict[str, Any]) -> bytes:
    """Render the structured critique to a PDF bytestream."""
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "TitleX", parent=base["Title"], fontSize=18, spaceAfter=10,
        ),
        "rec": ParagraphStyle(
            "Rec", parent=base["BodyText"], fontSize=11, spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontSize=12, spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "Body", parent=base["BodyText"], fontSize=10, alignment=TA_LEFT,
            leading=13,
        ),
    }
    rec_key = str(critique.get("recommendation") or "major_revision")
    rec_label = _REC_LABELS.get(rec_key, rec_key)
    rec_hex = _REC_COLORS_HEX.get(rec_key, "#374151")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        title=f"Peer Review — {source_title}",
    )
    story: list[Any] = [
        Paragraph("AI Peer Review", styles["title"]),
        Paragraph(_html_escape(source_title), styles["h2"]),
        Paragraph(
            f"<b>Recommendation:</b> <font color='{rec_hex}'><b>{_html_escape(rec_label)}</b></font>",
            styles["rec"],
        ),
        Spacer(1, 6),
    ]

    impression = (critique.get("overall_impression") or "").strip()
    if impression:
        story.append(Paragraph("Overall Impression", styles["h2"]))
        story.append(Paragraph(_html_escape(impression), styles["body"]))
        story.append(Spacer(1, 6))

    for key, label in _SECTIONS:
        items = critique.get(key) or []
        if not isinstance(items, list) or not items:
            continue
        story.append(Paragraph(label, styles["h2"]))
        bullets = [
            ListItem(Paragraph(_html_escape(str(it)), styles["body"]))
            for it in items
        ]
        story.append(
            ListFlowable(bullets, bulletType="bullet", start="circle", leftIndent=18)
        )
        story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


def render_critique_docx(
    *, source_title: str, critique: dict[str, Any]
) -> bytes:
    """Render the structured critique to a DOCX bytestream."""
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("AI Peer Review")
    run.bold = True
    run.font.size = Pt(18)

    src = doc.add_paragraph()
    sr = src.add_run(source_title)
    sr.italic = True
    sr.font.size = Pt(12)

    rec_key = str(critique.get("recommendation") or "major_revision")
    rec_label = _REC_LABELS.get(rec_key, rec_key)
    rec_color = _REC_COLORS_DOCX.get(rec_key, RGBColor(0x37, 0x41, 0x51))
    rec_para = doc.add_paragraph()
    label_run = rec_para.add_run("Recommendation: ")
    label_run.bold = True
    rec_run = rec_para.add_run(rec_label)
    rec_run.bold = True
    rec_run.font.color.rgb = rec_color

    impression = (critique.get("overall_impression") or "").strip()
    if impression:
        h = doc.add_paragraph()
        hr = h.add_run("Overall Impression")
        hr.bold = True
        hr.font.size = Pt(13)
        doc.add_paragraph(impression)

    for key, label in _SECTIONS:
        items = critique.get(key) or []
        if not isinstance(items, list) or not items:
            continue
        h = doc.add_paragraph()
        hr = h.add_run(label)
        hr.bold = True
        hr.font.size = Pt(13)
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
