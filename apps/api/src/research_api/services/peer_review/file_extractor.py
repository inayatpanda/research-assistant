"""Phase 4.6 — Extract text from an uploaded peer-review document.

Accepts PDF + DOCX. Magic-byte sniff (consistent with ``services/pdf_text``),
25 MiB hard cap, and an "all pages" extraction for PDFs (the existing
``extract_first_pages_text`` only reads the first ``n`` pages — useful for
citation extraction but the peer-review pipeline wants the entire document).
"""
from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader

from ..pdf_text import DOCX_MIME, PDF_MIME, detect_mime

MAX_UPLOAD_BYTES = 25 * 1024 * 1024  # 25 MiB hard cap


class PeerReviewExtractError(Exception):
    """Raised when extraction fails or input is rejected."""

    def __init__(self, message: str, *, status_code: int = 415) -> None:
        super().__init__(message)
        self.status_code = status_code


@dataclass
class FileExtraction:
    text: str
    mime: str
    source_type: str  # "uploaded_pdf" | "uploaded_docx"


def _from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts).strip()
    except Exception as exc:
        raise PeerReviewExtractError(
            f"Could not parse PDF: {exc}", status_code=422
        ) from exc


def _from_docx(data: bytes) -> str:
    # Same zip-bomb guard pattern as pdf_text._safe_zip_total_size_ok.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            total = sum(info.file_size for info in z.infolist())
            if total > 200 * 1024 * 1024:
                raise PeerReviewExtractError(
                    "DOCX uncompressed payload exceeds safety cap",
                    status_code=413,
                )
    except zipfile.BadZipFile as exc:
        raise PeerReviewExtractError(
            f"DOCX is not a valid zip archive: {exc}", status_code=422
        ) from exc

    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise PeerReviewExtractError(
            f"Could not parse DOCX: {exc}", status_code=422
        ) from exc

    parts: list[str] = [p.text for p in doc.paragraphs if (p.text or "").strip()]
    # Also surface table-cell text — peer reviewers care about reported
    # numbers, and TipTap tables typically end up here.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    parts.append(text)
    return "\n".join(parts).strip()


def extract_uploaded_document(
    *, data: bytes, original_filename: str | None = None
) -> FileExtraction:
    """Validate, sniff, and extract text from an uploaded PDF/DOCX."""
    if not data:
        raise PeerReviewExtractError("Empty file", status_code=400)
    if len(data) > MAX_UPLOAD_BYTES:
        raise PeerReviewExtractError(
            f"File exceeds {MAX_UPLOAD_BYTES // (1024 * 1024)} MiB cap",
            status_code=413,
        )
    mime = detect_mime(data)
    if mime == PDF_MIME:
        text = _from_pdf(data)
        return FileExtraction(text=text, mime=mime, source_type="uploaded_pdf")
    if mime == DOCX_MIME:
        text = _from_docx(data)
        return FileExtraction(text=text, mime=mime, source_type="uploaded_docx")
    raise PeerReviewExtractError(
        f"Unsupported file type ({mime}); only PDF and DOCX are accepted",
        status_code=415,
    )
